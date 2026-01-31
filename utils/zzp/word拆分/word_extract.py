from docx import Document
from docx.shared import Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os
from pathlib import Path
import tempfile
import shutil
import zipfile
from xml.etree import ElementTree as ET
from collections import defaultdict
import copy


class WordProjectExtractor:
    """Word文档项目提取器"""

    def __init__(self):
        self.paragraph_image_map = {}
        self.doc_structure = []  # 存储文档结构（段落和表格的顺序）
        self.table_positions = {}  # 存储表格位置信息

    def is_heading(self, paragraph, level=None):
        """
        判断段落是否为标题，如果指定level则检查特定级别
        level: None表示任意标题，填数字表示特定级别
        """
        if not paragraph.style or not paragraph.style.name:  # 如果段落没有样式或样式名，则该段落不是标题
            return False

        style_name = paragraph.style.name.lower()

        # 如果指定了级别，检查特定级别
        if level is not None:
            heading_patterns = [
                f'heading {level}',
                f'标题 {level}',
                f'heading{level}',
                f'title{level}',
                f'标题{level}'
            ]
            return any(pattern in style_name for pattern in heading_patterns)

        # 否则检查是否为任意标题
        return any(f'heading {i}' in style_name or f'标题 {i}' in style_name or
                   f'heading{i}' in style_name or f'标题{i}' in style_name
                   for i in range(1, 10))

    def get_heading_level(self, paragraph):
        """获取标题的级别，如果不是标题返回0"""
        if not paragraph.style or not paragraph.style.name:  # 若传入的段落没有样式或者样式没有名称，则直接返回0，表明不是标题
            return 0

        style_name = paragraph.style.name.lower()

        for i in range(1, 10):  # 遍历可能的标题级别
            patterns = [
                f'heading {i}',
                f'标题 {i}',
                f'heading{i}',
                f'title{i}',
                f'标题{i}'
            ]
            if any(pattern in style_name for pattern in patterns):
                return i

        return 0

    def get_outline_level(self, paragraph):
        """获取段落的大纲级别"""
        try:
            # 尝试从段落格式中获取大纲级别
            if hasattr(paragraph.paragraph_format, 'outline_level') and paragraph.paragraph_format.outline_level:
                return paragraph.paragraph_format.outline_level

            # 尝试从XML中直接提取大纲级别
            if hasattr(paragraph, '_element') and paragraph._element is not None:
                outline_elem = paragraph._element.find(
                    './/{http://schemas.openxmlformats.org/wordprocessingml/2006/main}outlineLvl')
                if outline_elem is not None and outline_elem.get(qn('w:val')):
                    return int(outline_elem.get(qn('w:val')))

            # 根据标题级别推断大纲级别
            heading_level = self.get_heading_level(paragraph)
            if heading_level > 0:
                return heading_level - 1  # 大纲级别通常比标题级别小1

            return 0  # 默认为正文级别
        except:
            return 0

    def set_outline_level(self, paragraph, level):
        """设置段落的大纲级别"""
        try:
            # 通过XML直接设置大纲级别
            if hasattr(paragraph, '_element') and paragraph._element is not None:
                # 查找或创建outlineLvl元素
                outline_elem = paragraph._element.find(  # 在段落的底层 XML 中查找名为 outlineLvl 的元素，该元素为段落的大纲级别
                    './/{http://schemas.openxmlformats.org/wordprocessingml/2006/main}outlineLvl')
                if outline_elem is None:  # 如果段落没有大纲级别，则创建一个新的大纲级别元素
                    outline_elem = OxmlElement('w:outlineLvl')
                    pPr = paragraph._element.find(
                        './/{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pPr')
                    if pPr is None:
                        pPr = OxmlElement('w:pPr')
                        paragraph._element.insert(0, pPr)  # 将大纲级别元素添加到段落属性元素中，从而为段落设置大纲级别
                    pPr.append(outline_elem)

                # 设置大纲级别
                outline_elem.set(qn('w:val'), str(level))
        except Exception as e:
            print(f"设置大纲级别失败: {e}")

    def apply_heading_style_by_outline_level(self, paragraph):
        """
        根据大纲级别应用对应的标题样式
        大纲级别1-9对应标题1-标题9样式
        """
        try:
            outline_level = self.get_outline_level(paragraph)
            if 1 <= outline_level <= 9:
                # 根据大纲级别设置对应的标题样式
                style_name = f"标题{outline_level}"
                # 尝试应用样式
                try:
                    paragraph.style = style_name
                    # print(f"   - 应用样式: {style_name} (大纲级别: {outline_level})")
                except Exception as e:
                    # 如果样式不存在，尝试使用英文样式名
                    try:
                        style_name_en = f"Heading {outline_level}"
                        paragraph.style = style_name_en
                        # print(f"   - 应用样式: {style_name_en} (大纲级别: {outline_level})")
                    except Exception as e2:
                        print(f"   - 警告: 无法应用样式 {style_name} 或 {style_name_en}，错误: {e2}")
        except Exception as e:
            print(f"   - 应用标题样式失败: {e}")

    def add_numbering_to_heading(self, paragraph, numbering_text, original_text=None):
        """
        为标题段落添加编号
        参数:
            paragraph: 要添加编号的段落
            numbering_text: 编号文本 (如 "1.1 - ")
            original_text: 原始文本内容（如果提供则使用，否则从段落中提取）
        """
        try:
            # 保存原始文本和格式
            original_runs = []
            if original_text is not None:
                # 如果提供了原始文本，创建一个临时的run信息
                original_runs.append({
                    'text': original_text,
                    'bold': False,
                    'italic': False,
                    'underline': False,
                    'font_size': None,
                    'font_name': None,
                    'font_color': None
                })
            else:
                # 否则从段落中提取
                for run in paragraph.runs:
                    original_runs.append({
                        'text': run.text,
                        'bold': run.bold,
                        'italic': run.italic,
                        'underline': run.underline,
                        'font_size': run.font.size,
                        'font_name': run.font.name,
                        'font_color': run.font.color.rgb if run.font.color.rgb else None
                    })

            # 清空段落
            paragraph.clear()

            # 添加编号
            numbering_run = paragraph.add_run(numbering_text)
            numbering_run.bold = False

            # 添加原始内容并保持格式
            for run_info in original_runs:
                new_run = paragraph.add_run(run_info['text'])
                new_run.bold = run_info['bold']
                new_run.italic = run_info['italic']
                new_run.underline = run_info['underline']
                if run_info['font_size']:
                    new_run.font.size = run_info['font_size']
                if run_info['font_name']:
                    new_run.font.name = run_info['font_name']
                if run_info['font_color']:
                    new_run.font.color.rgb = run_info['font_color']

        except Exception as e:
            print(f"添加编号失败: {e}")

    def generate_numbering_system_for_project(self, project_boundary, source_doc):
        """
        为单个项目生成独立的编号系统，每个项目都从1开始重新编号
        参数:
            project_boundary: 项目边界信息
            source_doc: 源文档对象
        返回:
            numbering_map: 编号映射字典
        """
        numbering_map = {}
        start, end = project_boundary['start'], project_boundary['end']

        # 获取项目标题的大纲级别作为基准
        project_title_para = source_doc.paragraphs[start]
        base_outline_level = self.get_outline_level(project_title_para)

        # 初始化各级别的计数器
        level_counters = {base_outline_level: 1}  # 项目标题从1开始
        current_levels = {base_outline_level: 1}

        # 为项目标题添加编号
        numbering_map[start] = {
            'numbering': "1 - ",
            'level': base_outline_level
        }

        # 遍历项目内除项目标题外的所有段落
        for para_idx in range(start + 1, end + 1):
            if para_idx >= len(source_doc.paragraphs):
                break

            # 获取该段落的大纲级别
            para = source_doc.paragraphs[para_idx]
            outline_level = self.get_outline_level(para)

            # 只处理标题级别的段落（大纲级别 > 0），即跳过正文，正文无需附加标号
            if outline_level > 0:
                level_diff = outline_level - base_outline_level

                if level_diff >= 0:
                    # 重置比当前级别低的所有计数器
                    for level in list(current_levels.keys()):    # 如果上一段落的标号为‘1.1.2’,而下一段落的标号，应为'1.2'，那么这里会现将'1.1.2'中的2删去
                        if level > outline_level:
                            del current_levels[level]

                    # 更新当前级别的计数器
                    if outline_level in current_levels:
                        current_levels[outline_level] += 1    #如果次级标题已经存在current_levels中，则该标题计数+1，如从‘1.1’——>'1.2'
                    else:
                        current_levels[outline_level] = 1       #如果次级标题之前不在current_levels中，则加入并赋值为1

                    # 构建编号
                    numbering_parts = []
                    for level in sorted(current_levels.keys()):
                        if level <= outline_level:
                            numbering_parts.append(str(current_levels[level]))

                    numbering_text = '.'.join(numbering_parts) + ' - '
                    numbering_map[para_idx] = {
                        'numbering': numbering_text,
                        'level': outline_level
                    }

        return numbering_map

    def extract_docx_images(self, input_path, output_dir):
        """
        从docx文件中提取所有图片到指定目录
        参数：
            input_path：输入的 word文档路径
            output_dir：提取的图片将被保存到的目录路径
        返回：
            image_files：列表，提取的图片保存到的路径
        """
        if not os.path.exists(output_dir):
            os.makedirs(output_dir)

        image_files = []

        try:
            # 解压docx文件
            with zipfile.ZipFile(input_path, 'r') as zip_ref:  # word本质是一个压缩包，使用zipfile.ZipFile打开
                # 提取所有图片文件
                for file_info in zip_ref.filelist:
                    if file_info.filename.startswith('word/media/'):  # 检查文件路径是否以 word/media/ 开头，这是word中存储图片的目录
                        filename = os.path.basename(file_info.filename)
                        if filename.lower().endswith(
                                ('.png', '.jpg', '.jpeg', '.gif', '.bmp')):  # 提取文件名，并检查扩展名是否为常见的图片格式
                            # 提取图片
                            zip_ref.extract(file_info.filename, output_dir)
                            src_path = os.path.join(output_dir, file_info.filename)
                            dst_path = os.path.join(output_dir, filename)

                            # 移动文件到目标目录
                            if os.path.exists(src_path):
                                if os.path.exists(dst_path):
                                    # 如果文件已存在，添加前缀
                                    base, ext = os.path.splitext(filename)
                                    dst_path = os.path.join(output_dir, f"{base}_{len(image_files)}{ext}")

                                shutil.move(src_path, dst_path)
                                image_files.append(dst_path)

                # 清理临时目录结构
                media_dir = os.path.join(output_dir, 'word', 'media')
                if os.path.exists(media_dir):
                    shutil.rmtree(os.path.join(output_dir, 'word'))

            print(f"   从文档中提取了 {len(image_files)} 张图片")
            return image_files

        except Exception as e:
            print(f"  ️  图片提取警告: {e}")
            return []

    def get_image_relationships(self, input_path):
        """
        获取文档中的图片关系映射
        参数：
            input_path：输入的Word文档路径
        返回：
            relationships：字典，存储图片关系映射
        """
        relationships = {}

        try:
            # 创建临时目录
            temp_dir = tempfile.mkdtemp()

            # 解压docx文件
            with zipfile.ZipFile(input_path, 'r') as zip_ref:
                # 提取关系文件
                rels_path = 'word/_rels/document.xml.rels'
                if rels_path in zip_ref.namelist():
                    zip_ref.extract(rels_path, temp_dir)

                    # 解析关系文件
                    rels_file = os.path.join(temp_dir, rels_path)
                    if os.path.exists(rels_file):
                        tree = ET.parse(rels_file)
                        root = tree.getroot()

                        for rel in root:  # 遍历关系文件的根元素root中的所有子元素rel
                            rel_id = rel.get('Id')
                            rel_type = rel.get('Type')
                            rel_target = rel.get('Target')  # 关系的目标路径

                            if 'image' in rel_type or 'media' in rel_target:  # 找出关系类型为image或目标路径为media的图片id和路径
                                relationships[rel_id] = rel_target

            # 清理临时目录
            shutil.rmtree(temp_dir)

        except Exception as e:
            print(f"️   获取图片关系失败: {e}")

        return relationships

    def find_precise_image_mapping(self, input_path):
        """
        为文档中的每个段落找到与之关联的图片
        参数：
            input_path：word路径
        返回：
            paragraph_images：字典，包含段落与图片映射关系，键为段落id，值为列表，包含与该段落关联的图片文件名
        """
        paragraph_images = defaultdict(list)

        try:
            # 创建临时目录
            temp_dir = tempfile.mkdtemp()

            # 使用 zipfile.ZipFile 打开输入的 Word 文档，并将其解压到临时目录中。
            with zipfile.ZipFile(input_path, 'r') as zip_ref:
                zip_ref.extractall(temp_dir)

            # 读取主文档XML
            document_path = os.path.join(temp_dir, 'word', 'document.xml')
            if not os.path.exists(document_path):  # 若不存在主文档XML文件，则返回空的字典
                return paragraph_images

            # 解析XML
            tree = ET.parse(document_path)
            root = tree.getroot()

            # 获取命名空间
            ns = {
                'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
                'a': 'http://schemas.openxmlformats.org/drawingml/2006/main',
                'r': 'http://schemas.openxmlformats.org/officeDocument/2006/relationships',
                'wp': 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing',
                'pic': 'http://schemas.openxmlformats.org/drawingml/2006/picture'
            }

            # 查找文档主体（<w:body>），并获取其中的所有段落（<w:p> 元素）
            body = root.find('w:body', ns)
            if body is None:
                return paragraph_images

            paragraphs = body.findall('w:p', ns)

            # 获取图片关系
            image_rels = self.get_image_relationships(input_path)  # 字典，键是word中图片的id，值为关系路径

            print(f"   找到 {len(image_rels)} 个图片关系")

            # 为每个段落查找图片
            for para_idx, para in enumerate(paragraphs):  # 遍历文档中的每个段落，para_idx 是段落的索引，para 是段落的 XML 元素
                # 方法1：查找blip元素（最常见的图片引用）
                blips = para.findall('.//a:blip', ns)  # 查找段落中的所有 <a:blip> 元素，这些元素通常用于引用图片
                for blip in blips:
                    embed_id = blip.get(
                        '{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')  # 获取 embed 属性的值，该值是图片的关系 ID
                    if embed_id and embed_id in image_rels:  # 如果该id存在于图片关系映射中，则获取对应的图片文件路径，并将其添加到当前段落的图片列表中。
                        image_file = image_rels[embed_id]
                        # 提取文件名
                        if '/' in image_file:
                            image_file = image_file.split('/')[-1]
                        paragraph_images[para_idx].append(image_file)

                # 方法2：查找inline中的图片
                drawings = para.findall('.//wp:inline', ns)
                for drawing in drawings:
                    blips = drawing.findall('.//a:blip', ns)
                    for blip in blips:
                        embed_id = blip.get(
                            '{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
                        if embed_id and embed_id in image_rels:
                            image_file = image_rels[embed_id]
                            if '/' in image_file:
                                image_file = image_file.split('/')[-1]
                            if image_file not in paragraph_images[para_idx]:
                                paragraph_images[para_idx].append(image_file)

                # 方法3：查找anchor中的图片
                anchors = para.findall('.//wp:anchor', ns)
                for anchor in anchors:
                    blips = anchor.findall('.//a:blip', ns)
                    for blip in blips:
                        embed_id = blip.get(
                            '{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
                        if embed_id and embed_id in image_rels:
                            image_file = image_rels[embed_id]
                            if '/' in image_file:
                                image_file = image_file.split('/')[-1]
                            if image_file not in paragraph_images[para_idx]:
                                paragraph_images[para_idx].append(image_file)

            # 清理临时目录
            shutil.rmtree(temp_dir)

        except Exception as e:
            print(f"   ️  建立图片映射失败: {e}")
            if 'temp_dir' in locals() and os.path.exists(temp_dir):
                shutil.rmtree(temp_dir)

        return paragraph_images

    def analyze_document_structure(self, doc):
        """
        分析文档结构，包括段落和表格的顺序
        返回项目标题层级、项目边界和文档结构
        参数：
            doc:传入的Document对象
        返回：
            project_level：识别到的项目标题层级
            project_boundaries：项目的边界信息（每个项目内容所处的段落索引范围）
        """
        paragraphs = doc.paragraphs  # 获取文档中所有的段落
        tables = doc.tables  # 获取文档中所有的表格

        # 获取文档的XML结构
        try:
            body = doc._element.body  # 获取文档的 XML 结构中的 <body> 元素
            self.doc_structure = []
            self.table_positions = {}

            # 遍历所有子元素
            for elem_idx, elem in enumerate(body):
                if elem.tag.endswith('p'):  # 如果元素以‘p’结尾，则代表段落
                    # 找到对应的段落索引
                    for i, para in enumerate(paragraphs):
                        if para._element == elem:
                            self.doc_structure.append({  # 存储文档结构信息，这里为段落
                                'type': 'paragraph',
                                'index': i,
                                'element': para,
                                'global_index': elem_idx
                            })
                            break
                elif elem.tag.endswith('tbl'):  # 表格元素
                    # 找到对应的表格索引，并记录到文档结构doc_structure和表格位置table_positions中
                    for i, table in enumerate(tables):
                        if table._element == elem:
                            # 记录表格位置
                            self.table_positions[i] = {  # 存储表格的位置信息，表格前的段落索引，用于判断表格属于哪一标题下
                                'table': table,
                                'global_index': elem_idx,
                                'prev_paragraph': self.find_previous_paragraph(elem_idx)
                            }

                            self.doc_structure.append({  # 存储文档结构信息，这里为表格的位置
                                'type': 'table',
                                'index': i,
                                'element': table,
                                'global_index': elem_idx
                            })
                            break

            print(f"文档结构分析完成: {len([x for x in self.doc_structure if x['type'] == 'paragraph'])} 个段落, "
                  f"{len([x for x in self.doc_structure if x['type'] == 'table'])} 个表格")
        except Exception as e:
            print(f"   ️ 文档结构分析失败: {e}")
            # 如果分析失败，使用简单的段落结构
            self.doc_structure = [{'type': 'paragraph', 'index': i, 'element': para}
                                  for i, para in enumerate(paragraphs)]

        # 统计各级标题的数量和位置
        heading_counts = defaultdict(int)
        heading_positions = defaultdict(list)  # 记录每个标题级别中所有标题的位置（段落索引）

        for i, para in enumerate(paragraphs):
            level = self.get_heading_level(para)  # 获取每个段落的标题级别
            if level > 0:  # 表明该段落是标题
                heading_counts[level] += 1
                heading_positions[level].append(i)

        print("文档标题统计:")
        for level in sorted(heading_counts.keys()):
            print(f"  {level}级标题: {heading_counts[level]} 个")

        # 智能识别项目标题层级
        project_level = None

        # 策略1: 如果一级标题只有一个，二级标题有多个，则项目在二级
        if heading_counts.get(1, 0) == 1 and heading_counts.get(2, 0) > 1:
            project_level = 2
            print("  识别项目标题层级: 2级标题")

        # 策略2: 如果一级标题有多个，则项目在一级
        elif heading_counts.get(1, 0) > 1:
            project_level = 1
            print("  识别项目标题层级: 1级标题")

        # 策略3: 如果二级标题只有一个，但三级标题有多个，则项目在三级
        elif heading_counts.get(2, 0) == 1 and heading_counts.get(3, 0) > 1:
            project_level = 3
            print("  识别项目标题层级: 3级标题")

        # 策略4: 针对只有一个项目的情况，默认使用最低的有标题的层级
        else:
            for level in range(1, 10):
                if heading_counts.get(level, 0) > 0:
                    project_level = level
                    print(f"  识别项目标题层级: {level}级标题 (默认)")
                    break

        if project_level is None:
            print("  未找到合适的项目标题层级")
            return None, []

        # 找到项目边界 - 从项目标题开始到下一个同级标题之前
        project_boundaries = []
        project_positions = heading_positions.get(project_level, [])  # 获取当前项目标题层级的所有标题位置

        for i, pos in enumerate(project_positions):
            start_pos = pos
            # 下一个同级标题的位置减1（如果存在下一个同级标题），否则是文档末尾。
            if i < len(project_positions) - 1:
                end_pos = project_positions[i + 1] - 1
            else:
                end_pos = len(paragraphs) - 1

            project_title = paragraphs[start_pos].text.strip()  # 当前项目标题的文本内容（即项目所在的标题）
            project_boundaries.append({
                'start': start_pos,
                'end': end_pos,
                'title': project_title,
                'level': project_level
            })

        return project_level, project_boundaries

    def find_previous_paragraph(self, table_global_index):
        """查找表格前最近的段落索引，确定表格与段落之间的关系时"""
        try:
            for i in range(table_global_index - 1, -1, -1):  # 从给定的表格位置向前遍历文档结构
                if i < len(self.doc_structure) and self.doc_structure[i]['type'] == 'paragraph':  # 查找第一个类型为段落的元素
                    return self.doc_structure[i]['index']
        except:
            pass
        return -1

    def copy_paragraph_format(self, source_para, target_para):
        """
        复制段落格式（不复制内容）
        参数：
            source_para：源段落对象
            target_para：目标段落对象
        """
        # 复制段落样式
        if source_para.style:
            target_para.style = source_para.style

        # 复制段落格式
        if source_para.paragraph_format:
            pf_source = source_para.paragraph_format  # 获取源段落的格式对象
            pf_target = target_para.paragraph_format

            if pf_source.alignment is not None:
                pf_target.alignment = pf_source.alignment  # 复制段落格式的对齐方式
            if pf_source.left_indent is not None:
                pf_target.left_indent = pf_source.left_indent  # 左缩进
            if pf_source.right_indent is not None:
                pf_target.right_indent = pf_source.right_indent  # 右缩进
            if pf_source.first_line_indent is not None:
                pf_target.first_line_indent = pf_source.first_line_indent  # 首行缩进
            if pf_source.space_before is not None:
                pf_target.space_before = pf_source.space_before  # 段前间距
            if pf_source.space_after is not None:
                pf_target.space_after = pf_source.space_after  # 段后间距

        # 复制大纲级别
        outline_level = self.get_outline_level(source_para)
        if outline_level > 0:
            self.set_outline_level(target_para, outline_level)

    def copy_run_format(self, source_run, target_run):
        """将源文本运行的格式（包括加粗、斜体、下划线、字体大小、字体名称和字体颜色）复制到目标文本运行"""
        target_run.bold = source_run.bold
        target_run.italic = source_run.italic
        target_run.underline = source_run.underline

        if source_run.font.size:
            target_run.font.size = source_run.font.size
        if source_run.font.name:
            target_run.font.name = source_run.font.name
        if source_run.font.color.rgb:
            target_run.font.color.rgb = source_run.font.color.rgb

    def clone_paragraph(self, doc, source_para, numbering_text=None):
        """
        克隆段落到新文档，保留格式和大纲级别，并根据大纲级别应用对应标题样式
        参数：
            doc: 目标文档对象，即新文档，段落将被克隆到这个文档中。
            source_para: 源文档中的段落对象，需要被克隆的段落。
            numbering_text: 可选的编号文本
        """
        # 在目标文档中建立新的段落对象，并将源段落的格式复制到新段落中
        new_para = doc.add_paragraph()
        self.copy_paragraph_format(source_para, new_para)

        # 复制文本和格式
        for run in source_para.runs:
            new_run = new_para.add_run(run.text)
            self.copy_run_format(run, new_run)

        # 如果提供了编号文本，添加到段落开头
        if numbering_text:
            self.add_numbering_to_heading(new_para, numbering_text)

        # 根据大纲级别应用对应的标题样式
        self.apply_heading_style_by_outline_level(new_para)

        return new_para

    def clone_table(self, doc, source_table):
        """克隆表格到新文档，完整复制包括表头在内的所有行和内容"""
        try:
            # 创建相同行列数的表格
            rows = len(source_table.rows)
            cols = len(source_table.columns)
            new_table = doc.add_table(rows=rows, cols=cols)

            # 复制表格样式
            if source_table.style:
                new_table.style = source_table.style

            # 复制表格格式（如列宽）
            for i, column in enumerate(source_table.columns):
                if i < len(new_table.columns):
                    new_table.columns[i].width = column.width

            # 复制所有行，包括表头
            for i, row in enumerate(source_table.rows):
                for j, cell in enumerate(row.cells):
                    if i < len(new_table.rows) and j < len(new_table.rows[i].cells):
                        new_cell = new_table.rows[i].cells[j]
                        # 清空默认内容
                        new_cell.text = ''

                        # 复制单元格的所有文本内容（简化版，不复制格式）
                        cell_text = ""
                        for para in cell.paragraphs:
                            for run in para.runs:
                                cell_text += run.text

                        # 直接设置单元格文本
                        new_cell.text = cell_text

            # 记录表头信息
            header_info = ""
            if rows > 0:
                header_cells = []
                for j in range(min(cols, 5)):  # 只取前5列作为表头示例
                    if j < len(source_table.rows[0].cells):
                        cell_text = source_table.rows[0].cells[j].text.strip()
                        if cell_text:
                            header_cells.append(cell_text[:20] + "..." if len(cell_text) > 20 else cell_text)
                if header_cells:
                    header_info = f", 表头: {' | '.join(header_cells)}"

            print(f"   - 复制表格: {rows} 行 × {cols} 列{header_info}")
            return new_table
        except Exception as e:
            print(f"   表格复制失败: {e}")
            # 如果详细复制失败，尝试简单复制
            try:
                # 创建表格并复制文本内容
                new_table = doc.add_table(rows=len(source_table.rows), cols=len(source_table.columns))
                for i, row in enumerate(source_table.rows):
                    for j, cell in enumerate(row.cells):
                        if i < len(new_table.rows) and j < len(new_table.rows[i].cells):
                            # 提取单元格所有文本
                            cell_text = ""
                            for para in cell.paragraphs:
                                for run in para.runs:
                                    cell_text += run.text
                            new_table.rows[i].cells[j].text = cell_text
                print(f"   - 简单复制表格: {len(source_table.rows)} 行 × {len(source_table.columns)} 列")
                return new_table
            except Exception as e2:
                print(f"   表格简单复制也失败: {e2}")
                return None

    def add_image_to_document(self, doc, image_path, width=Inches(6)):
        """
        向文档中添加图片
        参数：
            doc：目标文档
            image_path：图片路径
            width=Inches(6)：图片宽度，默认6英寸
        """
        try:
            if os.path.exists(image_path):
                doc.add_picture(image_path, width=width)
                return True
        except Exception as e:
            print(f"   ️  添加图片失败 {os.path.basename(image_path)}: {e}")
        return False

    def extract_projects(self, input_path, output_path):
        """
        提取符合查找条件的项目内容到新的word（一般情况，项目文档中的二级标题，一级为整个文档的标题）
        参数：
            input_path：传入的word路径
            output_path：存入到word路径
        返回：
            布尔值
        """
        try:
            source_doc = Document(input_path)
            target_doc = Document()  # 建了一个新的 Word 文档对象，用于存储提取的内容

            # 分析文档结构，智能识别项目标题层级和每个项目内容所处段落索引边界
            # project_level为项目所在的标题层级，project_boundaries为列表，元素为字典，包括项目开始的段落索引、结束的索引、标题内容（title）和标题层级
            project_level, project_boundaries = self.analyze_document_structure(source_doc)

            if project_level is None or not project_boundaries:
                print("   未找到任何项目")
                return False

            print(f"   找到 {len(project_boundaries)} 个项目")

            # 提取文档中的所有图片
            temp_image_dir = tempfile.mkdtemp()  # 创建一个临时目录
            all_images = self.extract_docx_images(input_path, temp_image_dir)  # 返回列表，包含提取的图片保存到的路径

            # 建立精确的图片映射关系
            self.paragraph_image_map = self.find_precise_image_mapping(input_path)  # 返回字典，键是段落的索引，值是与该段落关联的图片文件名列表

            # 复制保留的项目（为每一个项目创建一个word）
            for project_idx, boundary in enumerate(project_boundaries):
                start, end = boundary['start'], boundary['end']
                title = boundary['title']

                # 重新初始化目标文档
                target_doc = Document()

                # 为当前项目生成独立的编号系统（每个项目都从1开始）
                numbering_map = self.generate_numbering_system_for_project(boundary, source_doc)  # 字典，键为段落的索引，值为字典，包括标题的标号以及大纲级别

                # 添加项目标题，保持大纲级别，并添加编号
                title_numbering = numbering_map.get(start, {}).get('numbering', '1 - ')
                source_title_para = source_doc.paragraphs[start]

                # 创建新段落并复制格式
                title_para = target_doc.add_paragraph()
                self.copy_paragraph_format(source_title_para, title_para)

                # 添加编号和标题内容 - 修复：传递原始标题文本
                self.add_numbering_to_heading(title_para, title_numbering, title)

                # 对项目标题应用对应的标题样式
                self.apply_heading_style_by_outline_level(title_para)

                # 使用文档结构信息复制项目内容
                tables_added = 0
                paragraphs_added = 0

                # 遍历文档结构，复制在项目范围内的元素，分开处理段落（含文本和图片）与表格
                for elem_info in self.doc_structure:  # doc_structure字典，包含段落类型（type）、索引、段落内容（element）和word全局索引
                    if elem_info['type'] == 'paragraph':
                        para_idx = elem_info['index']
                        # 检查段落是否在项目范围内
                        if start < para_idx <= end:  # 从标题后开始
                            source_para = source_doc.paragraphs[para_idx]

                            # 检查是否需要添加编号
                            numbering_text = numbering_map.get(para_idx, {}).get('numbering')

                            # 克隆段落并添加编号（如果需要）
                            if numbering_text:
                                # 对于有编号的段落，使用clone_paragraph但不传入编号，因为我们要手动处理
                                new_para = target_doc.add_paragraph()
                                self.copy_paragraph_format(source_para, new_para)

                                # 复制文本内容
                                for run in source_para.runs:
                                    new_run = new_para.add_run(run.text)
                                    self.copy_run_format(run, new_run)

                                # 添加编号
                                self.add_numbering_to_heading(new_para, numbering_text)

                                # 应用标题样式
                                self.apply_heading_style_by_outline_level(new_para)
                            else:
                                # 对于没有编号的段落，正常克隆
                                self.clone_paragraph(target_doc, source_para)

                            paragraphs_added += 1

                            # 处理该段落中的图片
                            if para_idx in self.paragraph_image_map:
                                for image_file in self.paragraph_image_map[para_idx]:
                                    image_path = os.path.join(temp_image_dir, image_file)
                                    if os.path.exists(image_path):
                                        if self.add_image_to_document(target_doc, image_path):
                                            print(f"   - 添加图片: {image_file}")

                    elif elem_info['type'] == 'table':
                        table_idx = elem_info['index']
                        source_table = source_doc.tables[table_idx]

                        # 检查表格是否在项目范围内
                        # 通过查找表格前的段落来确定表格位置
                        if table_idx in self.table_positions:
                            prev_para_idx = self.table_positions[table_idx]['prev_paragraph']
                            # 如果表格前的段落在项目范围内，则表格属于该项目
                            if prev_para_idx != -1 and start <= prev_para_idx <= end:
                                cloned_table = self.clone_table(target_doc, source_table)
                                if cloned_table:
                                    tables_added += 1

                print(f"   - 为项目 '{title}' 添加了 {paragraphs_added} 个段落和 {tables_added} 个表格")

                # 保存文档
                output_dir = os.path.dirname(output_path) or "."
                if not os.path.exists(output_dir):
                    os.makedirs(output_dir, exist_ok=True)

                output_file_name = input_path[:-5] + f'_{project_idx + 1}.docx'
                target_doc.save(output_file_name)
                print(f"\n   提取完成。输出文件: {output_file_name}")

            # 清理临时文件
            shutil.rmtree(temp_image_dir)

            return True

        except Exception as e:
            print(f"   处理过程中出错: {e}")
            import traceback
            traceback.print_exc()

            # 清理临时文件
            if 'temp_image_dir' in locals() and os.path.exists(temp_image_dir):
                shutil.rmtree(temp_image_dir)

            return False

def extract_word_projects(input_path, output_path):
    """
    将传入的word中各项目内容做拆分，并单独存为一个word
    参数:
        input_path: 输入Word文档路径
        output_path: 输出Word文档路径
    """
    extractor = WordProjectExtractor()
    return extractor.extract_projects(input_path, output_path)

if __name__ == '__main__':
    # ================= 配置区域 =================
    # 在这里填入你要拆分的具体 Word 文档路径
    # 建议使用绝对路径，或者确保文件在当前目录下
    # 使用 r"" 可以防止路径中的反斜杠转义问题 (例如 r"C:\Users\Test\doc.docx")
    source_file_path = r"generate_report/utils/zzp/word拆分/附录X-1：信息系统建设与升级改造类（开发实施类）信息化项目可行性研究报告模板V6.0.docx" 
    
    # 设置输出目录 (留空 "" 则默认输出到脚本所在目录)
    output_directory = "./output"
    # ===========================================

    if os.path.exists(source_file_path):
        print(f"正在处理文件: {source_file_path} ...")
        
        # 确保输出目录存在
        if output_directory and not os.path.exists(output_directory):
            os.makedirs(output_directory)
            
        # 这里的 output_directory 主要用于指引目录，文件名会根据源文件名自动生成后缀
        # 注意：原类的逻辑中，生成的文件名是基于 input_path 修改的，
        # 如果你想强制保存到 output_directory，可能需要微调 extract_projects 中的 save 逻辑。
        # 但按照目前的代码，我们只需传入路径即可。
        
        success = extract_word_projects(source_file_path, output_directory)
        
        if success:
            print("拆分成功！")
        else:
            print("拆分过程中出现问题，或者未识别到项目。")
    else:
        print(f"错误: 找不到文件 '{source_file_path}'，请检查路径是否正确。")
