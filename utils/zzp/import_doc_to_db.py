import os
import sys
import pymysql
import traceback
import zipfile
import shutil
import tempfile
import xml.etree.ElementTree as ET
from collections import defaultdict
from sqlalchemy import create_engine, text
from urllib.parse import quote_plus, quote
from docx import Document
from docx.shared import Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from utils.zzp.docx_to_html import convert_docx_to_html

# ==========================================
# Monkey Patch for python-docx
# 修复 "There is no item named '.../NULL' in the archive" 错误
# ==========================================
import logging
from docx.opc.pkgreader import PackageReader
from docx.opc.package import Unmarshaller

# Patch 1: 忽略物理 ZIP 包中缺失的文件
def _patched_walk_phys_parts(phys_reader, srels, visited_partnames=None):
    if visited_partnames is None:
        visited_partnames = []

    for srel in srels:
        if srel.is_external:
            continue

        partname = srel.target_partname
        if partname in visited_partnames:
            continue

        visited_partnames.append(partname)

        try:
            blob = phys_reader.blob_for(partname)
        except KeyError:
            logging.warning(f"MonkeyPatch: Skipping missing part '{partname}' in docx archive (likely corrupted reference).")
            continue
        except Exception as e:
            logging.warning(f"MonkeyPatch: Error reading part '{partname}': {e}")
            continue

        reltype = srel.reltype
        part_srels = PackageReader._srels_for(phys_reader, partname)

        yield (partname, blob, reltype, part_srels)

        # 递归调用
        next_walker = PackageReader._walk_phys_parts(phys_reader, part_srels, visited_partnames)
        for partname, blob, reltype, srels in next_walker:
            yield (partname, blob, reltype, srels)

# Patch 2: 忽略指向缺失文件的关系引用
def _patched_unmarshal_relationships(pkg_reader, package, parts):
    """Add a relationship to the source object corresponding to each of the
    relationships in `pkg_reader` with its target_part set to the actual target part
    in `parts`."""
    for source_uri, srel in pkg_reader.iter_srels():
        # 1. 获取 Source
        if source_uri == "/":
            source = package
        else:
            if source_uri not in parts:
                logging.warning(f"MonkeyPatch: Skipping relationship for missing source '{source_uri}'")
                continue
            source = parts[source_uri]

        # 2. 获取 Target
        if srel.is_external:
            target = srel.target_ref
        else:
            if srel.target_partname not in parts:
                 logging.warning(f"MonkeyPatch: Skipping relationship to missing target '{srel.target_partname}' (rId: {srel.rId})")
                 continue
            target = parts[srel.target_partname]

        # 3. 加载关系
        source.load_rel(srel.reltype, target, srel.rId, srel.is_external)

# 应用 Patches
PackageReader._walk_phys_parts = staticmethod(_patched_walk_phys_parts)
Unmarshaller._unmarshal_relationships = staticmethod(_patched_unmarshal_relationships)

# ==========================================
# 0. 配置与环境
# ==========================================
project_root = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
if project_root not in sys.path:
    sys.path.append(project_root)

# 添加 generate_report_test 到 sys.path 以导入 server_config
generate_report_root = os.path.dirname(project_root)
if generate_report_root not in sys.path:
    sys.path.append(generate_report_root)
import server_config
from server_config import get_user_report_dir, ensure_user_directories

try:
    from zzp import sql_config as config
except ImportError:
    try:
        import sql_config as config
    except ImportError:
        print("无法导入 sql_config，请检查路径")
        sys.exit(1)

BASE_DIR = server_config.REPORT_DIR

# ==========================================
# 1. 数据库操作函数
# ==========================================

def get_db_connection():
    encoded_password = quote_plus(config.password)
    db_url = f"mysql+pymysql://{config.username}:{encoded_password}@{config.host}:{config.port}/{config.database}"
    return create_engine(db_url)

def get_or_create_report_type(conn, type_name):
    sql_check = text("SELECT id FROM report_type WHERE type_name = :name LIMIT 1")
    res = conn.execute(sql_check, {"name": type_name}).fetchone()
    if res:
        return res[0]
    sql_insert = text("INSERT INTO report_type (type_name) VALUES (:name)")
    res = conn.execute(sql_insert, {"name": type_name})
    return res.lastrowid

def insert_catalogue(conn, type_id, report_name_id, title, level, sort_order, parent_id, file_path):
    sql = text("""
        INSERT INTO report_catalogue 
        (type_id, report_name_id, catalogue_name, level, sortOrder, parent_id, file_name)
        VALUES 
        (:tid, :rid, :name, :lvl, :sort, :pid, :path)
    """)
    res = conn.execute(sql, {
        "tid": type_id,
        "rid": report_name_id,
        "name": title,
        "lvl": level,
        "sort": sort_order,
        "pid": parent_id,
        "path": file_path
    })
    return res.lastrowid

# ==========================================
# 2. 文档提取器类
# ==========================================

def get_heading_level(paragraph):
    """获取标题级别(1-9)，非标题返回0"""
    if not paragraph.style or not paragraph.style.name: return 0
    style_name = paragraph.style.name.lower()
    for i in range(1, 10):
        patterns = [f'heading {i}', f'标题 {i}', f'heading{i}', f'title{i}', f'标题{i}']
        if any(p in style_name for p in patterns): return i
    return 0

def scan_docx_structure(file_path):
    """
    快速扫描文档结构，不进行拆分。
    用于在文件上传后立即返回结构给前端。
    """
    try:
        doc = Document(file_path)
        structure = []
        current_level_counters = {}
        
        # 仅扫描段落寻找标题
        for para in doc.paragraphs:
            lvl = get_heading_level(para)
            if lvl > 0:
                title_text = para.text.strip()
                if not title_text: continue
                # [Fix] 忽略过长的“标题”，视其为正文误用样式
                if len(title_text) > 100:
                    continue

                # 维护编号逻辑 (为了与最终拆分结果一致)
                keys_to_del = [k for k in current_level_counters if k > lvl]
                for k in keys_to_del: del current_level_counters[k]
                current_level_counters[lvl] = current_level_counters.get(lvl, 0) + 1
                nums = [str(current_level_counters[k]) for k in sorted(current_level_counters.keys())]
                
                structure.append({
                    'title': title_text,
                    'level': lvl,
                    'numbering': ".".join(nums),
                    'sort_order': current_level_counters[lvl]
                })
        return structure
    except Exception as e:
        print(f"结构扫描失败: {e}")
        return []

class WordProjectExtractor:
    def __init__(self):
        self.paragraph_image_map = defaultdict(list)
        self.doc_structure = []  # 线性存储段落和表格

    def get_heading_level(self, paragraph):
        return get_heading_level(paragraph)


    # --- 图片提取与映射逻辑 ---

    def extract_docx_images(self, input_path, output_dir):
        """解压docx提取所有图片到临时目录"""
        if not os.path.exists(output_dir):
            os.makedirs(output_dir)
        
        image_files = []
        try:
            with zipfile.ZipFile(input_path, 'r') as zip_ref:
                for file_info in zip_ref.filelist:
                    if file_info.filename.startswith('word/media/'):
                        filename = os.path.basename(file_info.filename)
                        # 过滤常见图片格式
                        if filename.lower().endswith(('.png', '.jpg', '.jpeg', '.gif', '.bmp', '.emf', '.wmf')):
                            zip_ref.extract(file_info.filename, output_dir)
                            src_path = os.path.join(output_dir, file_info.filename)
                            dst_path = os.path.join(output_dir, filename)
                            
                            # 移动并扁平化路径
                            os.makedirs(os.path.dirname(dst_path), exist_ok=True)
                            shutil.move(src_path, dst_path)
                            image_files.append(dst_path)
            # 清理
            if os.path.exists(os.path.join(output_dir, 'word')):
                shutil.rmtree(os.path.join(output_dir, 'word'))
            return image_files
        except Exception as e:
            print(f"   [警告] 图片提取失败: {e}")
            return []

    def get_image_relationships(self, input_path):
        """解析 .rels 文件获取 rId 到图片路径的映射"""
        relationships = {}
        try:
            temp_dir = tempfile.mkdtemp()
            with zipfile.ZipFile(input_path, 'r') as zip_ref:
                rels_path = 'word/_rels/document.xml.rels'
                if rels_path in zip_ref.namelist():
                    zip_ref.extract(rels_path, temp_dir)
                    rels_file = os.path.join(temp_dir, rels_path)
                    tree = ET.parse(rels_file)
                    root = tree.getroot()
                    for rel in root:
                        rel_id = rel.get('Id')
                        rel_target = rel.get('Target')
                        if 'media' in rel_target:
                            relationships[rel_id] = rel_target
            shutil.rmtree(temp_dir)
        except Exception as e:
            print(f"   [警告] 关系解析失败: {e}")
        return relationships

    def find_precise_image_mapping(self, input_path):
        """建立 段落索引 -> [图片文件名列表] 的映射"""
        paragraph_images = defaultdict(list)
        try:
            temp_dir = tempfile.mkdtemp()
            with zipfile.ZipFile(input_path, 'r') as zip_ref:
                zip_ref.extract('word/document.xml', temp_dir)
            
            document_path = os.path.join(temp_dir, 'word', 'document.xml')
            tree = ET.parse(document_path)
            root = tree.getroot()
            
            ns = {
                'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main',
                'a': 'http://schemas.openxmlformats.org/drawingml/2006/main',
                'wp': 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing'
            }
            
            body = root.find('w:body', ns)
            paragraphs = body.findall('w:p', ns)
            image_rels = self.get_image_relationships(input_path)
            
            for para_idx, para in enumerate(paragraphs):
                # 查找所有可能的图片引用 blip
                blips = para.findall('.//a:blip', ns)
                for blip in blips:
                    embed_id = blip.get('{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed')
                    if embed_id and embed_id in image_rels:
                        image_file = image_rels[embed_id]
                        if '/' in image_file:
                            image_file = image_file.split('/')[-1]
                        paragraph_images[para_idx].append(image_file)
            
            shutil.rmtree(temp_dir)
        except Exception as e:
            print(f"   [警告] 映射建立失败: {e}")
            if os.path.exists(temp_dir): shutil.rmtree(temp_dir)
            
        return paragraph_images

    # --- 文档结构分析 ---

    def analyze_document_structure(self, doc):
        """将文档解析为扁平的结构列表（按顺序包含段落和表格）"""
        self.doc_structure = []
        body = doc._element.body
        paragraphs = doc.paragraphs
        tables = doc.tables
        
        for child in body:
            if child.tag.endswith('p'):
                for i, p in enumerate(paragraphs):
                    if p._element == child:
                        self.doc_structure.append({
                            'type': 'paragraph',
                            'index': i,
                            'obj': p
                        })
                        break
            elif child.tag.endswith('tbl'):
                for i, t in enumerate(tables):
                    if t._element == child:
                        self.doc_structure.append({
                            'type': 'table',
                            'index': i,
                            'obj': t
                        })
                        break

    # --- 复制功能 ---

    def copy_paragraph_format(self, src, tgt):
        if src.style: tgt.style = src.style
        if src.paragraph_format:
            try: tgt.paragraph_format.alignment = src.paragraph_format.alignment
            except: pass
            try: tgt.paragraph_format.first_line_indent = src.paragraph_format.first_line_indent
            except: pass

    def copy_run_format(self, src, tgt):
        tgt.bold = src.bold
        tgt.italic = src.italic
        tgt.underline = src.underline
        if src.font.size: tgt.font.size = src.font.size
        if src.font.color.rgb: tgt.font.color.rgb = src.font.color.rgb
        try: tgt.font.name = src.font.name
        except: pass

    def clone_paragraph_with_content(self, new_doc, src_para, numbering=None, image_dir=None, image_files=None):
        """克隆段落内容，包括文本、格式和图片"""
        p = new_doc.add_paragraph()
        self.copy_paragraph_format(src_para, p)
        
        # 1. 插入编号
        if numbering:
            nr = p.add_run(numbering)
            if src_para.runs:
                self.copy_run_format(src_para.runs[0], nr)
        
        # 2. 复制文本
        for r in src_para.runs:
            if r.text: 
                nr = p.add_run(r.text)
                self.copy_run_format(r, nr)
        
        # 3. 插入图片
        if image_files and image_dir:
            for img_name in image_files:
                img_path = os.path.join(image_dir, img_name)
                if os.path.exists(img_path):
                    try:
                        run = p.add_run()
                        run.add_picture(img_path, width=Inches(6.0))
                    except Exception as e:
                        print(f"      [图片] 插入失败 {img_name}: {e}")

    def clone_table(self, new_doc, src_table):
        """简单且稳健的表格复制"""
        try:
            rows = len(src_table.rows)
            cols = len(src_table.columns)
            new_table = new_doc.add_table(rows=rows, cols=cols)
            new_table.style = src_table.style
            
            for i, row in enumerate(src_table.rows):
                for j, cell in enumerate(row.cells):
                    new_table.rows[i].cells[j].text = cell.text.strip()
        except Exception as e:
            print(f"   [警告] 表格复制出错: {e}")

    # --- 主处理逻辑 ---

    def split_and_import_to_db(self, input_path, report_type_str, report_name_str, progress_callback=None, user_id=None):
        print(f"=== 开始处理: {report_name_str} (User: {user_id}) ===")
        
        # 0. 确保用户目录存在
        if user_id:
            ensure_user_directories(user_id)

        engine = get_db_connection()
        conn = engine.connect()
        trans = conn.begin()

        temp_image_dir = tempfile.mkdtemp()
        
        try:
            if progress_callback: progress_callback(5, "正在准备数据库记录...")

            # 1. 准备数据库
            type_id = get_or_create_report_type(conn, report_type_str)
            
            # 3. 检查重名 (增加user_id校验)
            if user_id is not None:
                # 针对特定用户的重名检查
                sql_check = text("SELECT id FROM report_name WHERE report_name = :name AND type_id = :tid AND user_id = :uid LIMIT 1")
                check_params = {"name": report_name_str, "tid": type_id, "uid": user_id}
            else:
                # 兼容旧逻辑（无用户ID），检查全局重名（或默认用户）
                # 这里假设无user_id时不应该与任何有user_id的冲突，或者检查 user_id IS NULL
                sql_check = text("SELECT id FROM report_name WHERE report_name = :name AND type_id = :tid AND user_id IS NULL LIMIT 1")
                check_params = {"name": report_name_str, "tid": type_id}
            
            existing_report = conn.execute(sql_check, check_params).fetchone()

            if existing_report:
                print(f"❌ 报告已存在: {report_name_str}")
                return False, f"在报告类型【{report_type_str}】下已存在名为【{report_name_str}】的报告"

            # 4. 插入报告记录
            if user_id is not None:
                result = conn.execute(text("INSERT INTO report_name (type_id, report_name, user_id) VALUES (:tid, :name, :uid)"), 
                            {"tid": type_id, "name": report_name_str, "uid": user_id})
            else:
                conn.execute(text("INSERT INTO report_name (type_id, report_name) VALUES (:tid, :name)"), 
                            {"tid": type_id, "name": report_name_str})
                            
            report_name_id = conn.execute(text("SELECT LAST_INSERT_ID()")).fetchone()[0]

            # 物理路径按用户隔离
            # 使用 server_config 中的路径配置
            user_report_base = get_user_report_dir(user_id)
            output_dir = os.path.join(user_report_base, report_type_str, report_name_str)
            
            if not os.path.exists(output_dir):
                os.makedirs(output_dir)

            # 2. 预处理：提取图片与建立映射
            print("正在提取图片资源...")
            if progress_callback: progress_callback(10, "正在提取图片资源...")
            self.extract_docx_images(input_path, temp_image_dir)
            self.paragraph_image_map = self.find_precise_image_mapping(input_path)
            
            # 3. 加载文档并分析结构
            if progress_callback: progress_callback(15, "正在分析文档深度结构...")
            source_doc = Document(input_path)
            self.analyze_document_structure(source_doc)

            # 4. 识别章节切分点
            paragraphs = source_doc.paragraphs
            current_level_counters = {}
            sections = []
            parent_id_stack = {0: 0}

            for i, para in enumerate(paragraphs):
                lvl = self.get_heading_level(para)
                if lvl > 0:
                    title_text = para.text.strip()
                    if not title_text: continue
                    
                    # [Fix] 忽略过长的“标题”，视其为正文误用样式，避免将正文误判为章节
                    if len(title_text) > 100:
                        continue

                    structure_index = -1
                    for idx, item in enumerate(self.doc_structure):
                        if item['type'] == 'paragraph' and item['index'] == i:
                            structure_index = idx
                            break
                    
                    if structure_index == -1: continue

                    keys_to_del = [k for k in current_level_counters if k > lvl]
                    for k in keys_to_del: del current_level_counters[k]
                    current_level_counters[lvl] = current_level_counters.get(lvl, 0) + 1
                    nums = [str(current_level_counters[k]) for k in sorted(current_level_counters.keys())]
                    
                    sections.append({
                        'title': title_text,
                        'level': lvl,
                        'numbering': ".".join(nums),
                        'sort_order': current_level_counters[lvl],
                        'structure_start_index': structure_index
                    })

            print(f"识别到 {len(sections)} 个章节，开始切分...")
            if progress_callback: progress_callback(20, f"识别到 {len(sections)} 个章节，准备开始切分与生成...")

            # 5. 执行切分与生成
            total_sections = len(sections)
            for idx, section in enumerate(sections):
                # 计算并更新进度
                current_percent = 20 + int((idx / total_sections) * 75) # 20% -> 95%
                if progress_callback and idx % 2 == 0: # 减少回调频率，每2个章节更新一次
                     progress_callback(current_percent, f"正在处理第 {idx+1}/{total_sections} 章: {section['title'][:15]}...")
                start_idx = section['structure_start_index']
                end_idx = sections[idx+1]['structure_start_index'] if idx + 1 < len(sections) else len(self.doc_structure)
                
                safe_title = section['title'].replace('/', '_').replace('\\', '_').replace(':', '_')[:50]
                file_name = f"{section['numbering']} {safe_title}.docx"
                file_path = os.path.join(output_dir, file_name)
                
                new_doc = Document()
                
                for i in range(start_idx, end_idx):
                    item = self.doc_structure[i]
                    
                    if item['type'] == 'paragraph':
                        src_para = item['obj']
                        para_index = item['index']
                        
                        num_str = f"{section['numbering']} " if i == start_idx else None
                        img_files = self.paragraph_image_map.get(para_index, [])
                        
                        self.clone_paragraph_with_content(
                            new_doc, 
                            src_para, 
                            numbering=num_str,
                            image_dir=temp_image_dir,
                            image_files=img_files
                        )
                        
                    elif item['type'] == 'table':
                        self.clone_table(new_doc, item['obj'])

                try:
                    new_doc.save(file_path)
                    print(f"   生成: {file_name}")
                    
                    # [新增] 生成 HTML
                    # 降低日志频率：仅当是该章节的最后一个文件时，或每隔一定数量时提示
                    # if progress_callback: progress_callback(current_percent, f"正在生成HTML: {safe_title[:15]}...")
                    
                    # 构造图片输出目录和URL前缀
                    images_dir = os.path.join(output_dir, "images")
                    if not os.path.exists(images_dir):
                        os.makedirs(images_dir)
                    
                    # URL 前缀: /python-api/report_files/{user_id}/{type}/{name}/images/
                    # 注意：report_files 映射到 REPORT_DIR，即 report/
                    # 路径结构: report/{user_id}/{type}/{name}/images/
                    url_prefix = f"/python-api/report_files/{user_id}/{quote(report_type_str)}/{quote(report_name_str)}/images/"
                    
                    convert_docx_to_html(
                        file_path, 
                        user_id=user_id,
                        image_output_dir=images_dir,
                        image_url_prefix=url_prefix
                    )
                    
                except Exception as e:
                    print(f"   [保存失败] {file_name}: {e}")

                current_lvl = section['level']
                parent_level = current_lvl - 1
                while parent_level > 0 and parent_level not in parent_id_stack:
                    parent_level -= 1
                parent_db_id = parent_id_stack.get(parent_level, 0)
                
                # [Fix] 数据库字段截断保护 (虽然前面有长度过滤，但这里做双重保险)
                # 同时，为了应对文档结构乱序导致的前端渲染问题，强制将章节编号写入标题
                raw_title = section['title']
                num_prefix = section['numbering']
                
                # 简单去重逻辑：如果原标题已经包含该编号，则不再拼接
                if raw_title.startswith(num_prefix + " ") or raw_title.startswith(num_prefix + "."):
                    full_title = raw_title
                else:
                    full_title = f"{num_prefix} {raw_title}"
                
                db_title = full_title[:250] if len(full_title) > 250 else full_title
                
                cat_id = insert_catalogue(
                    conn, type_id, report_name_id, db_title, 
                    current_lvl, section['sort_order'], parent_db_id, file_name
                )
                parent_id_stack[current_lvl] = cat_id

            trans.commit()
            if progress_callback: progress_callback(99, "所有章节处理完成，正在清理...")
            print("=== 处理完成 ===")
            return True, "导入成功"

        except Exception as e:
            trans.rollback()
            print("=== 异常回滚 ===")
            # traceback.print_exc() # 让外层捕获并记录日志
            raise e  # 抛出异常，让上层知道具体错误原因
        finally:
            conn.close()
            if os.path.exists(temp_image_dir):
                shutil.rmtree(temp_image_dir)

# ==========================================
# 4. 供路由调用的封装函数 (关键新增)
# ==========================================
def process_document(report_type, report_name, source_file, progress_callback=None, user_id=None):
    """
    路由调用的入口函数
    """
    if not os.path.exists(source_file):
        print(f"文件不存在: {source_file}")
        return False, f"文件不存在: {source_file}"
    
    extractor = WordProjectExtractor()
    return extractor.split_and_import_to_db(source_file, report_type, report_name, progress_callback, user_id=user_id)

# ==========================================
# 5. 测试入口
# ==========================================

if __name__ == "__main__":
    REPORT_TYPE = "可研究性报告1"
    REPORT_NAME = "1111" 
    SOURCE_FILE = r"/root/zzp/langextract-main/generate_report/utils/zzp/word拆分/附录X-1：信息系统建设与升级改造类（开发实施类）信息化项目可行性研究报告模板V6.0.docx"
    
    process_document(REPORT_TYPE, REPORT_NAME, SOURCE_FILE)