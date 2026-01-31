import os
import sys
import pymysql
from sqlalchemy import create_engine, text
from urllib.parse import quote_plus
from docx import Document
from collections import defaultdict

# ==========================================
# 0. 配置与环境
# ==========================================
project_root = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
if project_root not in sys.path:
    sys.path.append(project_root)

# 添加 generate_report_test 到 sys.path 以导入 server_config
generate_report_root = os.path.dirname(project_root)
if generate_report_root not in sys.path:
    sys.path.append(generate_report_root)
import server_config

try:
    from zzp import sql_config as config
except ImportError:
    try:
        import sql_config as config
    except ImportError:
        print("❌ 无法导入 sql_config，请检查路径")
        sys.exit(1)

# 全局基准路径
BASE_DIR = server_config.REPORT_DIR

# ==========================================
# 1. 数据库操作函数
# ==========================================

def get_db_connection():
    encoded_password = quote_plus(config.password)
    db_url = f"mysql+pymysql://{config.username}:{encoded_password}@{config.host}:{config.port}/{config.database}"
    return create_engine(db_url)

def get_or_create_report_type(conn, type_name):
    """获取或创建报告类型ID"""
    sql_check = text("SELECT id FROM report_type WHERE type_name = :name LIMIT 1")
    res = conn.execute(sql_check, {"name": type_name}).fetchone()
    if res:
        return res[0]
    sql_insert = text("INSERT INTO report_type (type_name) VALUES (:name)")
    res = conn.execute(sql_insert, {"name": type_name})
    return res.lastrowid

def insert_catalogue(conn, type_id, report_name_id, title, level, sort_order, parent_id, file_path):
    sql = text("""
        INSERT INTO report_catalogue 
        (type_id, report_name_id, catalogue_name, level, sortOrder, parent_id, file_name)
        VALUES 
        (:tid, :rid, :name, :lvl, :sort, :pid, :path)
    """)
    res = conn.execute(sql, {
        "tid": type_id,
        "rid": report_name_id,
        "name": title,
        "lvl": level,
        "sort": sort_order,
        "pid": parent_id,
        "path": file_path
    })
    return res.lastrowid

# ==========================================
# 2. 文档提取器类
# ==========================================

class WordProjectExtractor:
    def __init__(self):
        pass

    def get_heading_level(self, paragraph):
        """获取标题级别(1-9)，非标题返回0"""
        if not paragraph.style or not paragraph.style.name: return 0
        style_name = paragraph.style.name.lower()
        
        # 增加一些兼容性匹配
        for i in range(1, 10):
            patterns = [
                f'heading {i}', f'标题 {i}', 
                f'heading{i}', f'title{i}', f'标题{i}',
                f'level {i}', f'header {i}'
            ]
            if any(p in style_name for p in patterns): return i
        return 0

    def copy_paragraph_format(self, src, tgt):
        """复制段落格式"""
        if src.style: tgt.style = src.style
        if src.paragraph_format:
            try: tgt.paragraph_format.alignment = src.paragraph_format.alignment
            except: pass

    def copy_run_format(self, src, tgt):
        """复制字体格式"""
        tgt.bold = src.bold
        tgt.italic = src.italic
        if src.font.size: tgt.font.size = src.font.size
        if src.font.color.rgb: tgt.font.color.rgb = src.font.color.rgb

    def clone_paragraph(self, doc, source_para, numbering=None):
        """克隆段落到新文档"""
        p = doc.add_paragraph()
        self.copy_paragraph_format(source_para, p)
        if numbering:
            p.add_run(numbering)
        for r in source_para.runs:
            nr = p.add_run(r.text)
            self.copy_run_format(r, nr)
        return p

    def split_and_import_to_db(self, input_path, report_type_str, report_name_str):
        """
        读取文档 -> 查重 -> 线性切分 -> 保存文件 -> 写入数据库
        """
        print(f"=== 开始处理: {report_name_str} ===")
        
        # 1. 数据库连接与查重
        engine = get_db_connection()
        conn = engine.connect()
        trans = conn.begin() # 开启事务

        try:
            # 1.1 获取类型ID
            type_id = get_or_create_report_type(conn, report_type_str)

            # 1.2 报告名称查重
            sql_check_name = text("SELECT id FROM report_name WHERE report_name = :name AND type_id = :tid LIMIT 1")
            existing_report = conn.execute(sql_check_name, {"name": report_name_str, "tid": type_id}).fetchone()
            
            if existing_report:
                print(f"❌ 失败: 报告名称 '{report_name_str}' 已存在 (ID: {existing_report[0]})。")
                return False 

            # 1.3 插入新的报告名称
            sql_insert_name = text("INSERT INTO report_name (type_id, report_name) VALUES (:tid, :name)")
            res = conn.execute(sql_insert_name, {"tid": type_id, "name": report_name_str})
            report_name_id = res.lastrowid
            
            print(f"✅ 数据库准备就绪: TypeID={type_id}, New ReportID={report_name_id}")

            # 2. 准备物理目录
            output_dir = os.path.join(BASE_DIR, report_type_str, report_name_str)
            if not os.path.exists(output_dir):
                os.makedirs(output_dir)

            # 3. 加载文档
            source_doc = Document(input_path)
            paragraphs = source_doc.paragraphs
            
            # 4. 遍历与切分逻辑
            current_level_counters = {} 
            parent_id_stack = {0: 0} 
            sections = []
            
            # ======================================================
            # 第一遍扫描：确定有效标题和编号 (含调试打印)
            # ======================================================
            print("\n--- 正在分析文档结构 (打印前20个段落样式) ---")
            debug_count = 0 
            
            for i, para in enumerate(paragraphs):
                # >>> 调试代码：打印前 20 个非空段落的样式，帮你定位问题 <<<
                if debug_count < 20:
                    text_preview = para.text.strip()
                    if text_preview:
                        style_name = para.style.name if para.style else "无样式"
                        # 打印到控制台
                        print(f"🔍 [调试] 段落 {i} | 样式: '{style_name}' | 内容: '{text_preview[:15]}...'")
                        debug_count += 1
                # >>> 调试代码结束 <<<

                lvl = self.get_heading_level(para)
                if lvl > 0:
                    title_text = para.text.strip()
                    if not title_text:
                        continue # 跳过空标题

                    # 更新编号
                    keys_to_del = [k for k in current_level_counters if k > lvl]
                    for k in keys_to_del: del current_level_counters[k]
                    
                    current_level_counters[lvl] = current_level_counters.get(lvl, 0) + 1
                    
                    nums = [str(current_level_counters[k]) for k in sorted(current_level_counters.keys())]
                    numbering_str = ".".join(nums)
                    
                    sections.append({
                        'start_index': i,
                        'level': lvl,
                        'title': title_text,
                        'numbering': numbering_str,
                        'sort_order': current_level_counters[lvl]
                    })
            
            if not sections:
                print("\n⚠️  [严重] 文档中未识别到任何标题！")
                print("    请检查上方 [调试] 日志，确认该文档的标题是否使用了标准的 '标题 1' / 'Heading 1' 样式。")
                print("    如果样式名是自定义的（如 '公文一级标题'），请修改 get_heading_level 方法。")
                return False

            print(f"\n识别到 {len(sections)} 个有效章节，开始拆分生成...")

            # --- 第二遍扫描：生成文件并入库 ---
            for idx, section in enumerate(sections):
                start_idx = section['start_index']
                end_idx = sections[idx+1]['start_index'] if idx + 1 < len(sections) else len(paragraphs)
                
                # 构建文件名
                safe_title = section['title'].replace('/', '_').replace('\\', '_').replace(':', '_').replace('*', '_').replace('?', '_').replace('"', '_').replace('<', '_').replace('>', '_').replace('|', '_')
                file_name = f"{section['numbering']} {safe_title}.docx"
                file_path = os.path.join(output_dir, file_name)
                
                # 创建新文档
                new_doc = Document()
                
                # 复制标题
                src_title_para = paragraphs[start_idx]
                self.clone_paragraph(new_doc, src_title_para, numbering=f"{section['numbering']} ")
                
                # 复制正文
                for p_idx in range(start_idx + 1, end_idx):
                    src_p = paragraphs[p_idx]
                    self.clone_paragraph(new_doc, src_p)
                    
                # 保存文件
                new_doc.save(file_path)
                print(f"   生成: {file_name}")
                
                # --- 数据库入库 ---
                current_lvl = section['level']
                
                # 计算 Parent ID
                parent_level = current_lvl - 1
                while parent_level > 0 and parent_level not in parent_id_stack:
                    parent_level -= 1
                parent_db_id = parent_id_stack.get(parent_level, 0)
                
                # 插入记录
                new_catalogue_id = insert_catalogue(
                    conn, 
                    type_id, 
                    report_name_id, 
                    section['title'], 
                    current_lvl, 
                    section['sort_order'], 
                    parent_db_id, 
                    file_path
                )
                
                # 更新栈
                parent_id_stack[current_lvl] = new_catalogue_id

            trans.commit()
            print("=== ✅ 处理完成，数据已入库 ===")
            return True

        except Exception as e:
            trans.rollback()
            print(f"=== ❌ 处理失败: {e} ===")
            import traceback
            traceback.print_exc()
            return False
        finally:
            conn.close()

# ==========================================
# 3. 主入口
# ==========================================

def process_document(report_type, report_name, source_file):
    """
    执行切分入库流程
    Returns: bool (True 成功, False 失败)
    """
    if not os.path.exists(source_file):
        print(f"❌ 文件不存在: {source_file}")
        return False
    
    extractor = WordProjectExtractor()
    return extractor.split_and_import_to_db(source_file, report_type, report_name)

if __name__ == "__main__":
    # 配置参数（仅用于本地单独运行测试）
    REPORT_TYPE = "可研究性报告1"
    REPORT_NAME = "test_to_doc3"
    SOURCE_FILE = r"/root/zzp/langextract-main/generate_report/utils/zzp/word拆分/附录X-1：信息系统建设与升级改造类（开发实施类）信息化项目可行性研究报告模板V6.0.docx"
    
    print(f"开始切分文档...\n源文件: {SOURCE_FILE}")
    process_document(REPORT_TYPE, REPORT_NAME, SOURCE_FILE)