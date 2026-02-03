import os
import json
import shutil
import pymysql
import re
import unicodedata
import hashlib
from sqlalchemy import create_engine, text
from urllib.parse import quote_plus, quote
from docx import Document
import sys
from utils.zzp.docx_to_html import convert_docx_to_html
# ==========================================
# 文件名 / 路径安全处理（修复非法命名问题）
# ==========================================
WINDOWS_RESERVED_NAMES = {
    "CON", "PRN", "AUX", "NUL",
    *(f"COM{i}" for i in range(1, 10)),
    *(f"LPT{i}" for i in range(1, 10)),
}

def safe_filename(name: str, max_length: int = 150) -> str:
    """
    将任意用户输入转换为安全的文件名（不含扩展名）
    适配 Windows / Linux / macOS
    """
    if not name:
        return "untitled"

    # 1. Unicode 规范化
    name = unicodedata.normalize("NFKC", name)

    # 2. 移除控制字符
    name = re.sub(r"[\x00-\x1f\x7f]", "", name)

    # 3. 替换 Windows 非法字符
    name = re.sub(r'[\\/:*?"<>|]', "_", name)

    # 4. 合并多余空白
    name = re.sub(r"\s+", " ", name).strip()

    # 5. 去除结尾点和空格（Windows 会炸）
    name = name.rstrip(" .")

    # 6. 防止保留名
    if name.upper() in WINDOWS_RESERVED_NAMES:
        name = f"_{name}"

    # 7. 截断长度
    return name[:max_length]


def safe_path_component(name: str) -> str:
    """
    专用于目录名（reportType / reportName）
    """
    return safe_filename(name, max_length=100)


# ==========================================
# 0. 解决配置导入路径问题
# ==========================================
project_root = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
if project_root not in sys.path:
    sys.path.append(project_root)

# 添加 generate_report_test 到 sys.path 以导入 server_config
generate_report_root = os.path.dirname(project_root)
if generate_report_root not in sys.path:
    sys.path.append(generate_report_root)
import server_config

from zzp import sql_config as config

# ==========================================
# 1. 基础配置与数据库连接
# ==========================================

BASE_DIR = server_config.REPORT_DIR

def get_db_connection():
    encoded_password = quote_plus(config.password)
    db_url = f"mysql+pymysql://{config.username}:{encoded_password}@{config.host}:{config.port}/{config.database}?charset=utf8mb4"
    return create_engine(db_url, pool_recycle=3600, pool_pre_ping=True)

# ==========================================
# 2. 数据库查询与写入函数
# ==========================================

def get_source_file_path(connection, origin_type, origin_report, origin_title, user_id=None, origin_id=None):
    """
    根据来源信息，查询源文件的存储路径
    优先查找用户私有模板，其次查找公共模板
    [Best Practice] 如果提供 origin_id (Source ID)，则优先使用 ID 精确查找
    """
    
    # 策略 1: ID 精确查找 (Best Practice)
    if origin_id:
        sql_id = text("SELECT file_name FROM report_catalogue WHERE id = :oid")
        result_id = connection.execute(sql_id, {"oid": origin_id}).fetchone()
        if result_id and result_id[0]:
            db_path = result_id[0]
            if os.path.exists(db_path):
                return db_path
            # 如果 ID 对应的路径不存在，尝试智能推断 (同下文逻辑)
            # 但为简化逻辑，这里若 ID 查到的文件不存在，我们仍继续尝试下面的标题匹配兜底
            # 或者复用下文的路径推断逻辑。这里选择复用下文逻辑。
            # 为了复用，我们先暂时不返回，而是让代码继续流转? 
            # 不，ID 查到的文件名是最准确的，应该基于这个文件名去推断路径。
            
            # 复用路径推断逻辑
            file_name = os.path.basename(db_path)
            # 尝试路径 1: 用户私有目录
            if user_id:
                user_path = os.path.join(server_config.get_user_report_dir(user_id), origin_type, origin_report, file_name)
                if os.path.exists(user_path):
                    return user_path
            # 尝试路径 2: 公共目录
            public_path = os.path.join(server_config.get_user_report_dir(None), origin_type, origin_report, file_name)
            if os.path.exists(public_path):
                return public_path

    # 策略 2: 标题模糊匹配 (Legacy / Fallback)
    # [Fix] 增加标题清洗逻辑，解决因标题中包含旧编号导致的匹配失败问题
    clean_origin_title = re.sub(r'^[\d\.]+\s*', '', origin_title).strip() if origin_title else ""

    sql = text("""
        SELECT c.file_name, n.user_id
        FROM report_catalogue c
        JOIN report_name n ON c.report_name_id = n.id
        JOIN report_type t ON c.type_id = t.id
        WHERE t.type_name = :otype 
          AND n.report_name = :oname 
          AND (
              c.catalogue_name = :otitle 
              OR c.catalogue_name = :clean_otitle
          )
          AND (n.user_id = :uid OR n.user_id IS NULL)
        ORDER BY n.user_id DESC
        LIMIT 1
    """)
    result = connection.execute(sql, {
        "otype": origin_type,
        "oname": origin_report,
        "otitle": origin_title,
        "clean_otitle": clean_origin_title,
        "uid": user_id
    }).fetchone()
    
    if result and result[0]:
        db_path = result[0]
        
        # 1. 如果路径直接存在，完美
        if os.path.exists(db_path):
            return db_path
            
        # 2. 如果不存在，尝试智能推断路径 (可能是因为多用户隔离导致物理路径变更)
        # 提取文件名
        file_name = os.path.basename(db_path)
        
        # 尝试路径 1: 用户私有目录
        if user_id:
            user_path = os.path.join(server_config.get_user_report_dir(user_id), origin_type, origin_report, file_name)
            if os.path.exists(user_path):
                return user_path
                
        # 尝试路径 2: 公共目录
        public_path = os.path.join(server_config.get_user_report_dir(None), origin_type, origin_report, file_name)
        if os.path.exists(public_path):
            return public_path
            
        # 尝试路径 3: 数据库里的路径可能本身就是相对路径或者是旧的绝对路径，尝试拼接
        # 这里可以根据情况扩展
            
        return None
    return None

def get_or_create_report_type(connection, type_name, user_id=None):
    # 1. 尝试查找用户私有类型
    if user_id is not None:
        query_check = text("SELECT id FROM report_type WHERE type_name = :type_name AND user_id = :uid LIMIT 1")
        result = connection.execute(query_check, {"type_name": type_name, "uid": user_id}).fetchone()
        if result:
            return result[0]

    # 2. 尝试查找公共类型
    query_check = text("SELECT id FROM report_type WHERE type_name = :type_name AND user_id IS NULL LIMIT 1")
    result = connection.execute(query_check, {"type_name": type_name}).fetchone()
    if result:
        return result[0]
    
    # 3. 创建新类型 (有用户则私有，无用户则公共)
    if user_id is not None:
        insert_sql = text("INSERT INTO report_type (type_name, user_id) VALUES (:type_name, :uid)")
        result = connection.execute(insert_sql, {"type_name": type_name, "uid": user_id})
    else:
        insert_sql = text("INSERT INTO report_type (type_name) VALUES (:type_name)")
        result = connection.execute(insert_sql, {"type_name": type_name})
    return result.lastrowid

def get_or_create_report_name(connection, type_db_id, report_name, user_id=None):
    # 根据 user_id 区分报告，允许不同用户有同名报告
    if user_id is not None:
        query_check = text("""
            SELECT id FROM report_name 
            WHERE report_name = :report_name AND type_id = :type_id AND user_id = :user_id LIMIT 1
        """)
        params = {"report_name": report_name, "type_id": type_db_id, "user_id": user_id}
    else:
        # 兼容旧逻辑或无用户场景
        query_check = text("""
            SELECT id FROM report_name 
            WHERE report_name = :report_name AND type_id = :type_id LIMIT 1
        """)
        params = {"report_name": report_name, "type_id": type_db_id}

    result = connection.execute(query_check, params).fetchone()
    
    if result:
        return result[0]
    else:
        # 插入时带上 user_id
        if user_id is not None:
            insert_sql = text("INSERT INTO report_name (type_id, report_name, user_id) VALUES (:type_id, :report_name, :user_id)")
            result = connection.execute(insert_sql, {"type_id": type_db_id, "report_name": report_name, "user_id": user_id})
        else:
            insert_sql = text("INSERT INTO report_name (type_id, report_name) VALUES (:type_id, :report_name)")
            result = connection.execute(insert_sql, {"type_id": type_db_id, "report_name": report_name})
        return result.lastrowid

def insert_catalogue_item(connection, data):
    sql = text("""
        INSERT INTO report_catalogue 
        (type_id, report_name_id, catalogue_name, level, sortOrder, parent_id, file_name)
        VALUES 
        (:type_id, :report_name_id, :catalogue_name, :level, :sortOrder, :parent_id, :file_name)
    """)
    result = connection.execute(sql, data)
    return result.lastrowid

# ==========================================
# 3. 文件生成逻辑
# ==========================================

def generate_prefix(parent_prefix, sort_order):
    if parent_prefix:
        return f"{parent_prefix}.{sort_order}"
    else:
        return f"{sort_order}"

def create_docx_file(file_path, content_title):
    """创建新文件（当 isimport=0 时使用）"""
    doc = Document()
    doc.add_heading(content_title, level=0)
    doc.add_paragraph(f"这是新创建的章节【{content_title}】。")
    doc.save(file_path)

def process_node_recursive(connection, node, root_path, parent_prefix, parent_db_id, context_ids, created_files=None, user_id=None):
    if created_files is None:
        created_files = []
        
    type_db_id = context_ids['type_db_id']
    report_name_db_id = context_ids['report_name_db_id']
    
    # 获取节点基本信息
    title = node.get("title")
    level = node.get("level")
    sort_order = node.get("sortOrder")
    is_import = node.get("isimport", 0)  # 获取是否导入标记
    children = node.get("children", [])
    
    # 1. 生成文件名
    # [Fix] 修复文件名包含非法字符导致创建失败的问题
    safe_title = title.replace('/', '_').replace('\\', '_').replace(':', '_')
    current_prefix = generate_prefix(parent_prefix, sort_order)
    
    # [Action] 强力清洗标题中的旧编号
    clean_title = re.sub(r'^[\d\.]+\s*', '', title).strip()

    # 文件名组装
    raw_node_name = f"{current_prefix} {clean_title}"
    
    node_name = raw_node_name
    safe_node_name = safe_filename(raw_node_name)
    file_name_with_ext = f"{safe_node_name}.docx"
    target_file_path = os.path.join(root_path, file_name_with_ext)

    
    # 2. 处理文件生成逻辑 (核心修改部分)
    file_created = False
    
    if is_import == 1:
        # === 导入模式 ===
        origin_type = node.get("originreportType")
        origin_report = node.get("originreportName")
        origin_title = node.get("origintitle")
        origin_id = node.get("origin_catalogue_id") # [Best Practice] 尝试获取源 ID
        
        # A. 去数据库查找源文件路径
        source_path = get_source_file_path(connection, origin_type, origin_report, origin_title, user_id=user_id, origin_id=origin_id)
        
        # B. 执行复制
        if source_path and os.path.exists(source_path):
            try:
                # copy2 会保留文件的元数据（如修改时间）
                shutil.copy2(source_path, target_file_path)
                # print(f" -> [Copy] 从 {source_path} 复制到 {target_file_path}")
                file_created = True
            except Exception as e:
                print(f" -> [Error] 复制文件失败: {e}")
        else:
            print(f" -> [Warn] 源文件未找到: {origin_report}-{origin_title}, 将创建一个空文件代替。")
            # 如果源文件找不到，为了保证目录完整，我们降级为创建一个新文件
            create_docx_file(target_file_path, f"{node_name} (源文件缺失)")
            file_created = True
            
    else:
        # === 新建模式 (isimport=0) ===
        try:
            create_docx_file(target_file_path, node_name)
            file_created = True
        except Exception as e:
            print(f" -> [Error] 创建文件失败: {e}")

    # 3. 数据库入库 (记录新的文件路径)
    if file_created:
        # [修改] 不再在此处同步生成 HTML，而是收集路径后续异步处理
        created_files.append(target_file_path)

        db_data = {
            "type_id": type_db_id,
            "report_name_id": report_name_db_id,
            "catalogue_name": clean_title,  # [Fix] 存入数据库时使用清洗后的标题，避免幽灵编号残留
            "level": level,
            "sortOrder": sort_order,
            "parent_id": parent_db_id,
            "file_name": target_file_path 
        }
        current_node_db_id = insert_catalogue_item(connection, db_data)
    else:
        current_node_db_id = 0 # 理论上不应发生，除非写入失败
    
    # 4. 递归处理子节点
    for child in children:
        process_node_recursive(
            connection, 
            child, 
            root_path=root_path,
            parent_prefix=current_prefix, 
            parent_db_id=current_node_db_id, 
            context_ids=context_ids,
            created_files=created_files,
            user_id=user_id
        )
    return created_files

# ==========================================
# 4. 主入口函数 (已修改：增加查重逻辑)
# ==========================================

def generate_merged_report_from_json(json_data, agent_user_id=None):
    engine = get_db_connection()
    
    report_type_str = json_data.get("reportType")
    report_name_str = json_data.get("reportName")
    chapters = json_data.get("chapters", [])
    
    print(f"=== 开始处理合并报告: {report_name_str} ===")
    
    # ---------------------------------------------------------
    # [新增逻辑] 步骤 0: 预先检查报告名是否重复
    # ---------------------------------------------------------
    try:
        # 使用 connect() 进行只读查询，不开启事务
        with engine.connect() as conn:
            # 1. 先找 Type ID (优先找私有，再找公共)
            type_id = None
            if agent_user_id is not None:
                type_sql = text("SELECT id FROM report_type WHERE type_name = :t_name AND user_id = :uid LIMIT 1")
                type_res = conn.execute(type_sql, {"t_name": report_type_str, "uid": agent_user_id}).fetchone()
                if type_res:
                    type_id = type_res[0]
            
            if not type_id:
                type_sql = text("SELECT id FROM report_type WHERE type_name = :t_name AND user_id IS NULL LIMIT 1")
                type_res = conn.execute(type_sql, {"t_name": report_type_str}).fetchone()
                if type_res:
                    type_id = type_res[0]
            
            if type_id:
                # 2. 再查该 Type 下有没有重名的 Report
                # [MODIFIED] Check user_id if present
                if agent_user_id is not None:
                     check_name_sql = text("SELECT id FROM report_name WHERE report_name = :r_name AND type_id = :t_id AND user_id = :uid LIMIT 1")
                     params = {"r_name": report_name_str, "t_id": type_id, "uid": agent_user_id}
                else:
                     check_name_sql = text("SELECT id FROM report_name WHERE report_name = :r_name AND type_id = :t_id LIMIT 1")
                     params = {"r_name": report_name_str, "t_id": type_id}

                existing_report = conn.execute(check_name_sql, params).fetchone()
                
                if existing_report:
                    # [关键点] 如果查到了，直接抛出异常，中断后续操作
                    print(f"⚠️ 报告名称重复: {report_name_str}")
                    raise ValueError("DUPLICATE_REPORT_NAME")
                    
    except ValueError as ve:
        # 将 ValueError 继续向上抛出，给 Router 捕获
        raise ve
    except Exception as e:
        print(f"查重步骤出错: {e}")
        # 如果是数据库连接错误等，也抛出，避免强行创建
        raise e

    # ---------------------------------------------------------
    # 步骤 1: 准备根目录 (查重通过后才创建文件夹)
    # ---------------------------------------------------------
    # [MODIFIED] 使用用户隔离的路径
    base_dir = server_config.get_user_report_dir(agent_user_id)
    # 确保用户目录存在
    if agent_user_id is not None and not os.path.exists(base_dir):
        try:
            os.makedirs(base_dir)
        except OSError:
            pass
            
    safe_report_type = safe_path_component(report_type_str)
    safe_report_name = safe_path_component(report_name_str)

    root_path = os.path.join(base_dir, safe_report_type, safe_report_name)

    if not os.path.exists(root_path):
        os.makedirs(root_path)
    
    created_files = []
    try:
        # 开启事务进行写入
        with engine.begin() as connection:
            # 2. 准备基础信息
            type_db_id = get_or_create_report_type(connection, report_type_str, user_id=agent_user_id)
            # 注意：这里的 get_or_create 依然保留，作为双重保险，或者你可以改成纯 insert
            report_name_db_id = get_or_create_report_name(connection, type_db_id, report_name_str, user_id=agent_user_id)
            
            context_ids = {
                "type_db_id": type_db_id,
                "report_name_db_id": report_name_db_id
            }
            
            # 3. 递归处理
            for chapter in chapters:
                process_node_recursive(
                    connection, 
                    chapter, 
                    root_path=root_path, 
                    parent_prefix="", 
                    parent_db_id=0, 
                    context_ids=context_ids,
                    created_files=created_files,
                    user_id=agent_user_id
                )
                
        print("=== ✅ 报告合并及生成成功！ ===")
        return created_files
        
    except Exception as e:
        print(f"=== ❌ 处理失败: {e} ===")
        import traceback
        traceback.print_exc()
        # 再次抛出异常确保 Router 知道出错了
        raise e

def generate_html_for_report_background(file_paths):
    """
    后台任务：为生成的文件创建 HTML 版本
    """
    if not file_paths:
        return
    
    print(f"开始后台生成 HTML，共 {len(file_paths)} 个文件...")
    for path in file_paths:
        try:
            rel_path = os.path.relpath(path, server_config.REPORT_DIR)
            parts = rel_path.split(os.sep)
            user_id = None
            report_type = None
            report_name = None

            if len(parts) >= 4 and parts[0].isdigit():
                user_id = parts[0]
                report_type = parts[1]
                report_name = parts[2]
            elif len(parts) >= 3:
                report_type = parts[0]
                report_name = parts[1]

            if report_type and report_name:
                if user_id:
                    images_dir = os.path.join(
                        server_config.EDITOR_IMAGE_DIR,
                        "report",
                        str(user_id),
                        report_type,
                        report_name
                    )
                    url_prefix = f"/python-api/editor_images/report/{user_id}/{quote(report_type)}/{quote(report_name)}/"
                else:
                    images_dir = os.path.join(
                        server_config.EDITOR_IMAGE_DIR,
                        "report",
                        report_type,
                        report_name
                    )
                    url_prefix = f"/python-api/editor_images/report/{quote(report_type)}/{quote(report_name)}/"
                if not os.path.exists(images_dir):
                    os.makedirs(images_dir)
                convert_docx_to_html(
                    path,
                    user_id=int(user_id) if user_id else None,
                    image_output_dir=images_dir,
                    image_url_prefix=url_prefix
                )
            else:
                convert_docx_to_html(path)
        except Exception as e:
            print(f"后台生成 HTML 失败: {path}, 错误: {e}")
    print("后台 HTML 生成任务完成。")

# ==========================================
# 5. 本地测试运行
# ==========================================
if __name__ == "__main__":
    # 使用你提供的混合数据进行测试
    input_json = {
      "reportName": "可研究性报告（合并资产报告）",
      "reportType": "可研究性报告1",
      "chapters": [
        {
          "title": "申请人基本情况",
          "origintitle": "申请人基本情况",
          "originreportType": "资产报告",
          "originreportName": "通用资产报告",
          "isimport": 1,
          "level": 1,
          "sortOrder": 1,
          "children": [
            {
              "title": "申请人基本的情况",
              "origintitle": "申请人基本的情况",
              "originreportType": "资产报告",
              "originreportName": "通用资产报告",
              "isimport": 1,
              "level": 2,
              "sortOrder": 1,
              "children": []
            }
          ]
        },
        {
          "title": "申请人经营情况",
          "origintitle": "申请人经营情况",
          "originreportType": "资产报告",
          "originreportName": "通用资产报告",
          "isimport": 1,
          "level": 1,
          "sortOrder": 2,
          "children": [
            {
              "title": "生产经营状况",
              "origintitle": "生产经营状况",
              "originreportType": "资产报告",
              "originreportName": "通用资产报告",
              "isimport": 1,
              "level": 2,
              "sortOrder": 1,
                "children": [
                   {
                        "title": "经营分析",
                        "origintitle": "经营分析",
                        "originreportType": "资产报告",
                        "originreportName": "通用资产报告",
                        "isimport": 1,
                        "level": 3,
                        "sortOrder": 1,
                        "children": []
                    }
              ]
            }
          ]
        },    
        {
          "title": "风险分析",
          "origintitle": "",
          "originreportType": "",
          "originreportName": "",
          "isimport": 0,
          "level": 1,
          "sortOrder": 3,
          "children": []
        }
      ]
    }
    # 执行合并生成逻辑
    generate_merged_report_from_json(input_json)
    
